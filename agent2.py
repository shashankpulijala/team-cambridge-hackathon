from uagents import Agent, Bureau, Context, Model
from uagents.setup import fund_agent_if_low
from docx import Document
import openai
import pandas as pd
import re

rules_address = 'agent1qgut44dptpp9t65mtmjsms3ffj2m8d0e0df290nk0rd2celd7y84ssmevxw'

class Message(Model):
    message: str

class Response(Model):
    response: str


class ErrorMessage(Model):
    error: str

parser = Agent(name="parser", seed="jkhgfdsa", port=5001, endpoint=['http://localhost:5001/submit'])

fund_agent_if_low(parser.wallet.address())

# Declare a global variable
global_sender_add = None

@parser.on_event('startup')
async def user_handler(ctx: Context):
    ctx.logger.info(f'My address is {parser.address}')

@parser.on_query(model=Message)
async def parse_doc(ctx: Context, sender: str, msg: Message):
    global global_sender_add
    doc = Document(msg.message)
    global_sender_add = sender
    full_text = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(' | '.join(row_text))
    document_text = '\n'.join(full_text)

    ctx.logger.info('Converted table to text')

    openai.api_key = 'sk-proj-GRDjhCH5HEgVBCxMYmiVT3BlbkFJEpsCezMzvhKU13Nm6aIm'

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "Extract and map the key fields text for an insurance quote:"},
            {"role": "user", "content": document_text}
        ],
        max_tokens=1500,
        temperature=0.5,
    )
    mapped_fields = response.choices[0].message['content'].strip()
    ctx.logger.info(mapped_fields)

    output_file = 'mapped_fields.json'

    with open(output_file, 'w') as file:
        file.write(mapped_fields)

    ctx.logger.info('Generated mapping file')

    await ctx.send(rules_address, Message(message=mapped_fields))

@parser.on_message(model=Response)
async def handle_response(ctx: Context, sender: str, msg: Response):
    global global_sender_add
    ctx.logger.info(f'Got response for query {msg.response}')
    await ctx.send(global_sender_add, Message(message=msg.response))

if __name__ == "__main__":
    parser.run()
